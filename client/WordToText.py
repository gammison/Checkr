import docx

doc = docx.Document('/Users/Latter/Documents/UVa Problems.docx')
paragraph_length = len(doc.paragraphs)
text_file = open('text/uva.txt', 'w')
for para in doc.paragraphs:
    text_file.write(para.text+'\n')

